# 🤖 BERT-ENHANCED ML SERVICE
# Integrates BERT/Transformers with existing resume analysis structure

import os
import sys
import joblib
import pickle
import numpy as np
import pandas as pd
from pathlib import Path
import traceback
import fitz  # PyMuPDF
from docx import Document
import io
import re
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')

# Add current directory to path
current_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(current_dir)

# Import BERT and transformer libraries
try:
    from sentence_transformers import SentenceTransformer, util
    BERT_AVAILABLE = True
    print("🤖 SentenceTransformers available")
except ImportError as e:
    print(f"⚠️ SentenceTransformers not available: {e}")
    BERT_AVAILABLE = False
    # Create a dummy util for fallback
    class DummyUtil:
        @staticmethod
        def cos_sim(a, b):
            return [[0.5]]
    util = DummyUtil()

# Fallback to traditional ML
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics.pairwise import cosine_similarity

class BERTEnhancedMLService:
    """BERT-Enhanced ML service for advanced resume analysis"""
    
    def __init__(self, models_path=None, use_bert=True):
        self.models_path = models_path or os.path.join(current_dir, 'trained_models')
        self.models = {}
        self.vectorizers = {}
        self.encoders = {}
        self.is_loaded = False
        self.use_bert = use_bert and BERT_AVAILABLE
        
        # Initialize BERT models if available
        if self.use_bert:
            self.init_bert_models()
        
        # Load traditional models as fallback
        self.load_traditional_models()
        
        # Initialize skill ontology
        self.init_skill_ontology()
    
    def init_bert_models(self):
        """Initialize SentenceTransformer models"""
        if not BERT_AVAILABLE:
            print("⚠️ BERT not available, using traditional methods only")
            self.sentence_model = None
            self.bert_initialized = False
            return
            
        try:
            print("🤖 Initializing BERT models...")
            
            # Load SentenceTransformer for semantic embeddings
            self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
            print("✅ SentenceTransformer loaded")
            
            # Precompute job category embeddings for fast classification
            self.precompute_category_embeddings()
            
            self.bert_initialized = True
            
        except Exception as e:
            print(f"⚠️ BERT initialization failed: {e}")
            self.sentence_model = None
            self.use_bert = False
            self.bert_initialized = False
    
    def precompute_category_embeddings(self):
        """Precompute embeddings for job categories"""
        self.job_categories = [
            "Software Engineering", "Data Science", "Machine Learning", "Frontend Development",
            "Backend Development", "Full Stack Development", "DevOps Engineering", 
            "Cloud Engineering", "Mobile Development", "Product Management",
            "UI/UX Design", "Quality Assurance", "Database Administration",
            "Network Engineering", "Cybersecurity", "Business Analysis",
            "Project Management", "Technical Writing", "Sales Engineering"
        ]
        
        if hasattr(self, 'sentence_model') and self.sentence_model:
            self.category_embeddings = self.sentence_model.encode(self.job_categories)
            print(f"✅ Precomputed embeddings for {len(self.job_categories)} job categories")
        else:
            self.category_embeddings = None
            print("⚠️ SentenceTransformer not available, skipping embeddings precomputation")
    
    def load_traditional_models(self):
        """Load traditional ML models as fallback"""
        try:
            print(f"🔧 Loading traditional models from: {self.models_path}")
            
            # Core ML models
            model_files = {
                'category_classifier': 'category_classifier.pkl',
                'experience_predictor': 'experience_predictor.pkl',
                'match_score_predictor': 'match_score_predictor.pkl',
                'category_tfidf': 'category_tfidf.pkl',
                'experience_tfidf': 'experience_tfidf.pkl',
                'match_score_tfidf': 'match_score_tfidf.pkl'
            }
            
            # Load models if they exist
            for model_name, filename in model_files.items():
                filepath = os.path.join(self.models_path, filename)
                if os.path.exists(filepath):
                    try:
                        self.models[model_name] = joblib.load(filepath)
                        print(f"✅ Loaded {model_name}")
                    except Exception as e:
                        print(f"⚠️ Failed to load {model_name}: {e}")
                        self.models[model_name] = None
                else:
                    print(f"⚠️ Model file not found: {filepath}")
                    self.models[model_name] = None
            
            # Load training stats if available
            stats_file = os.path.join(self.models_path, 'training_stats.json')
            if os.path.exists(stats_file):
                import json
                with open(stats_file, 'r') as f:
                    self.training_stats = json.load(f)
                print("✅ Loaded training statistics")
            else:
                self.training_stats = {}
            
            # Check if we have at least basic models
            essential_models = ['category_classifier', 'category_tfidf']
            if all(self.models.get(model) is not None for model in essential_models):
                self.is_loaded = True
                print("✅ Traditional models loaded successfully!")
            else:
                print("⚠️ Some traditional models missing, creating fallback")
                self.create_fallback_models()
                
        except Exception as e:
            print(f"❌ Error loading traditional models: {e}")
            self.create_fallback_models()
    
    def create_fallback_models(self):
        """Create simple fallback models if trained models aren't available"""
        print("🔧 Creating fallback models...")
        
        try:
            # Create simple models with dummy data
            sample_texts = [
                "Python developer with machine learning experience",
                "Java backend engineer with spring boot",
                "React frontend developer with JavaScript",
                "Data scientist with pandas and numpy",
                "DevOps engineer with AWS and Docker"
            ]
            
            sample_categories = [
                "Software Engineering", "Backend Development", "Frontend Development", 
                "Data Science", "DevOps Engineering"
            ]
            
            # TF-IDF Vectorizer
            self.models['category_tfidf'] = TfidfVectorizer(max_features=1000, stop_words='english')
            X_sample = self.models['category_tfidf'].fit_transform(sample_texts)
            
            # Label Encoder
            le = LabelEncoder()
            y_sample = le.fit_transform(sample_categories)
            
            # Category Classifier
            self.models['category_classifier'] = RandomForestClassifier(n_estimators=50, random_state=42)
            self.models['category_classifier'].fit(X_sample, y_sample)
            
            # Experience Predictor
            self.models['experience_predictor'] = RandomForestRegressor(n_estimators=50, random_state=42)
            sample_experience = [3, 5, 2, 4, 6]
            self.models['experience_predictor'].fit(X_sample, sample_experience)
            
            # Store label encoder
            self.encoders['category_le'] = le
            
            # Copy TF-IDF for other models
            self.models['experience_tfidf'] = self.models['category_tfidf']
            self.models['match_score_tfidf'] = self.models['category_tfidf']
            self.models['match_score_predictor'] = self.models['experience_predictor']
            
            self.is_loaded = True
            print("✅ Fallback models created successfully!")
            
        except Exception as e:
            print(f"❌ Failed to create fallback models: {e}")
            self.is_loaded = False
    
    def init_skill_ontology(self):
        """Initialize comprehensive skill ontology"""
        self.skill_ontology = {
            'programming_languages': [
                'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'Go', 'Rust', 
                'Ruby', 'PHP', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'Perl'
            ],
            'web_technologies': [
                'React', 'Angular', 'Vue.js', 'Node.js', 'Express', 'Django', 'Flask', 
                'Spring Boot', 'Laravel', 'Ruby on Rails', 'ASP.NET', 'Next.js'
            ],
            'databases': [
                'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Cassandra', 'Neo4j', 
                'Oracle', 'SQL Server', 'SQLite', 'DynamoDB'
            ],
            'cloud_platforms': [
                'AWS', 'Azure', 'Google Cloud', 'IBM Cloud', 'DigitalOcean', 
                'Heroku', 'Vercel', 'Netlify', 'Firebase'
            ],
            'data_science': [
                'Machine Learning', 'Deep Learning', 'Data Science', 'AI', 'NLP', 
                'Computer Vision', 'TensorFlow', 'PyTorch', 'Keras', 'Scikit-learn',
                'Pandas', 'NumPy', 'Matplotlib', 'Seaborn'
            ]
        }
        
        # Flatten all skills for easy access
        self.all_skills = []
        for category, skills in self.skill_ontology.items():
            self.all_skills.extend(skills)
    
    def parse_document(self, file_content, filename):
        """Parse PDF or DOCX document with enhanced text extraction"""
        try:
            file_ext = Path(filename).suffix.lower()
            
            if file_ext == '.pdf':
                return self.parse_pdf_enhanced(file_content)
            elif file_ext in ['.docx', '.doc']:
                return self.parse_docx_enhanced(file_content)
            elif file_ext == '.txt':
                return file_content.decode('utf-8')
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            print(f"❌ Document parsing error: {e}")
            return ""
    
    def parse_pdf_enhanced(self, file_content):
        """Enhanced PDF parsing with better text extraction"""
        try:
            # Try PyMuPDF first (better for complex layouts)
            doc = fitz.open(stream=file_content, filetype="pdf")
            text = ""
            
            for page in doc:
                text += page.get_text()
            doc.close()
            
            # Clean up text
            text = self.clean_extracted_text(text)
            return text
            
        except Exception as e:
            print(f"⚠️ PyMuPDF failed, trying fallback: {e}")
            # Fallback to simpler extraction
            return ""
    
    def parse_docx_enhanced(self, file_content):
        """Enhanced DOCX parsing"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Extract text from paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # Clean up text
            text = self.clean_extracted_text(text)
            return text
            
        except Exception as e:
            print(f"❌ DOCX parsing error: {e}")
            return ""
    
    def clean_extracted_text(self, text):
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\-\(\)\@\+\#]', ' ', text)
        
        # Normalize spacing
        text = ' '.join(text.split())
        
        return text.strip()
    
    def analyze_resume_bert(self, resume_text):
        """BERT-enhanced resume analysis"""
        if not self.use_bert or not hasattr(self, 'sentence_model'):
            return self.analyze_resume_traditional(resume_text)
        
        try:
            print("🤖 Running BERT-enhanced analysis...")
            
            # Get semantic embeddings
            resume_embedding = self.sentence_model.encode([resume_text])
            
            # BERT-based category classification
            category_result = self.classify_category_bert(resume_text, resume_embedding)
            
            # BERT-enhanced skill extraction
            skills_result = self.extract_skills_bert(resume_text)
            
            # BERT-based experience prediction
            experience_result = self.predict_experience_bert(resume_text, resume_embedding)
            
            # Semantic job matching
            match_scores = self.calculate_semantic_match_scores(resume_embedding)
            
            # Combine traditional and BERT results
            traditional_result = self.analyze_resume_traditional(resume_text)
            
            # Enhanced result with BERT insights
            result = {
                **traditional_result,
                'bert_enhanced': True,
                'semantic_category': category_result,
                'semantic_skills': skills_result,
                'semantic_experience': experience_result,
                'semantic_matches': match_scores,
                'confidence_scores': {
                    'category': category_result.get('confidence', 0.0),
                    'skills': skills_result.get('confidence', 0.0),
                    'experience': experience_result.get('confidence', 0.0)
                }
            }
            
            # Override traditional results with BERT if confidence is high
            if category_result.get('confidence', 0) > 0.8:
                result['predicted_category'] = category_result['category']
                result['category_confidence'] = category_result['confidence']
            
            if experience_result.get('confidence', 0) > 0.7:
                result['predicted_experience'] = experience_result['experience']
                result['experience_confidence'] = experience_result['confidence']
            
            print("✅ BERT-enhanced analysis completed")
            return result
            
        except Exception as e:
            print(f"⚠️ BERT analysis failed, falling back to traditional: {e}")
            return self.analyze_resume_traditional(resume_text)
    
    def classify_category_bert(self, resume_text, resume_embedding):
        """BERT-based job category classification"""
        try:
            # Calculate semantic similarity with job categories
            similarities = util.cos_sim(resume_embedding, self.category_embeddings)[0]
            
            # Get best match
            best_idx = similarities.argmax().item()
            best_score = similarities[best_idx].item()
            
            return {
                'category': self.job_categories[best_idx],
                'confidence': float(best_score),
                'top_matches': [
                    {
                        'category': self.job_categories[i],
                        'score': float(similarities[i])
                    }
                    for i in similarities.argsort(descending=True)[:5]
                ]
            }
            
        except Exception as e:
            print(f"⚠️ BERT category classification error: {e}")
            return {'category': 'Unknown', 'confidence': 0.0}
    
    def extract_skills_bert(self, resume_text):
        """BERT-enhanced skill extraction"""
        try:
            # Traditional regex-based extraction
            extracted_skills = self.extract_skills_traditional(resume_text)
            
            if not hasattr(self, 'sentence_model'):
                return {'skills': extracted_skills, 'confidence': 0.5}
            
            # Semantic skill matching using BERT
            semantic_skills = []
            
            # Get embeddings for resume text
            resume_embedding = self.sentence_model.encode([resume_text])
            
            # Get embeddings for all skills
            skill_embeddings = self.sentence_model.encode(self.all_skills)
            
            # Calculate similarities
            similarities = util.cos_sim(resume_embedding, skill_embeddings)[0]
            
            # Filter skills with high semantic similarity
            for i, skill in enumerate(self.all_skills):
                if similarities[i] > 0.4:  # Threshold for semantic match
                    semantic_skills.append({
                        'skill': skill,
                        'confidence': float(similarities[i]),
                        'source': 'semantic'
                    })
            
            # Combine traditional and semantic skills
            all_skills = list(set(extracted_skills + [s['skill'] for s in semantic_skills]))
            
            return {
                'skills': all_skills,
                'semantic_skills': semantic_skills,
                'traditional_skills': extracted_skills,
                'confidence': 0.8,
                'total_found': len(all_skills)
            }
            
        except Exception as e:
            print(f"⚠️ BERT skill extraction error: {e}")
            return {'skills': self.extract_skills_traditional(resume_text), 'confidence': 0.5}
    
    def predict_experience_bert(self, resume_text, resume_embedding):
        """BERT-enhanced experience prediction"""
        try:
            # Traditional experience prediction
            traditional_exp = self.predict_experience_traditional(resume_text)
            
            # BERT-based experience indicators
            experience_keywords = {
                'junior': ['intern', 'trainee', 'entry', 'junior', 'associate', 'fresher'],
                'mid': ['developer', 'engineer', 'analyst', 'specialist', 'coordinator'],
                'senior': ['senior', 'lead', 'principal', 'staff', 'manager', 'architect'],
                'executive': ['director', 'vp', 'cto', 'ceo', 'head', 'chief']
            }
            
            # Count experience indicators in text
            text_lower = resume_text.lower()
            experience_scores = {}
            
            for level, keywords in experience_keywords.items():
                score = sum(1 for keyword in keywords if keyword in text_lower)
                experience_scores[level] = score
            
            # Predict based on highest score
            if max(experience_scores.values()) > 0:
                predicted_level = max(experience_scores, key=experience_scores.get)
                level_mapping = {'junior': 1, 'mid': 3, 'senior': 7, 'executive': 12}
                bert_experience = level_mapping.get(predicted_level, traditional_exp)
                confidence = min(0.9, max(experience_scores.values()) / 3)
            else:
                bert_experience = traditional_exp
                confidence = 0.6
            
            return {
                'experience': max(1, min(15, int(bert_experience))),
                'confidence': confidence,
                'level': predicted_level if 'predicted_level' in locals() else 'mid',
                'indicators': experience_scores
            }
            
        except Exception as e:
            print(f"⚠️ BERT experience prediction error: {e}")
            return {'experience': 3, 'confidence': 0.5}
    
    def calculate_semantic_match_scores(self, resume_embedding):
        """Calculate semantic job matching scores"""
        try:
            # Sample job descriptions for matching
            job_descriptions = [
                "Python developer with Django and machine learning experience",
                "Frontend developer with React and JavaScript skills",
                "Data scientist with pandas, numpy, and machine learning",
                "DevOps engineer with AWS, Docker, and Kubernetes",
                "Full stack developer with Python and React"
            ]
            
            job_embeddings = self.sentence_model.encode(job_descriptions)
            similarities = util.cos_sim(resume_embedding, job_embeddings)[0]
            
            matches = []
            for i, desc in enumerate(job_descriptions):
                matches.append({
                    'job_description': desc,
                    'similarity_score': float(similarities[i]),
                    'match_percentage': float(similarities[i] * 100)
                })
            
            # Sort by similarity
            matches.sort(key=lambda x: x['similarity_score'], reverse=True)
            
            return matches[:3]  # Return top 3 matches
            
        except Exception as e:
            print(f"⚠️ Semantic matching error: {e}")
            return []
    
    def analyze_resume_traditional(self, resume_text):
        """Traditional ML-based resume analysis (fallback)"""
        if not self.is_loaded:
            return self.get_fallback_analysis(resume_text)
        
        try:
            # Extract features using TF-IDF
            if self.models.get('category_tfidf'):
                text_features = self.models['category_tfidf'].transform([resume_text])
            else:
                return self.get_fallback_analysis(resume_text)
            
            # Predict category
            if self.models.get('category_classifier'):
                category_pred = self.models['category_classifier'].predict(text_features)[0]
                if hasattr(self.encoders.get('category_le'), 'inverse_transform'):
                    category = self.encoders['category_le'].inverse_transform([category_pred])[0]
                else:
                    category = f"Category_{category_pred}"
            else:
                category = "Software Engineering"
            
            # Predict experience
            if self.models.get('experience_predictor'):
                experience_pred = self.models['experience_predictor'].predict(text_features)[0]
                experience = max(1, min(15, int(experience_pred)))
            else:
                experience = self.predict_experience_traditional(resume_text)
            
            # Extract skills
            skills = self.extract_skills_traditional(resume_text)
            
            # Calculate match score
            if self.models.get('match_score_predictor'):
                match_score = self.models['match_score_predictor'].predict(text_features)[0]
                match_score = max(0, min(100, match_score * 100))
            else:
                match_score = 75.0
            
            return {
                'predicted_category': category,
                'predicted_experience': experience,
                'extracted_skills': skills,
                'match_score': match_score,
                'analysis_method': 'traditional_ml',
                'model_confidence': 0.85,
                'bert_enhanced': False
            }
            
        except Exception as e:
            print(f"⚠️ Traditional analysis error: {e}")
            return self.get_fallback_analysis(resume_text)
    
    def extract_skills_traditional(self, text):
        """Traditional skill extraction using regex"""
        skills_found = []
        text_lower = text.lower()
        
        for skill in self.all_skills:
            if skill.lower() in text_lower:
                skills_found.append(skill)
        
        return list(set(skills_found))
    
    def predict_experience_traditional(self, text):
        """Traditional experience prediction using keywords"""
        # Simple keyword-based experience estimation
        text_lower = text.lower()
        
        # Look for year mentions
        year_pattern = r'(\d+)\s*(?:years?|yrs?)'
        years_found = re.findall(year_pattern, text_lower)
        
        if years_found:
            return max(int(year) for year in years_found)
        
        # Fallback based on title keywords
        if any(keyword in text_lower for keyword in ['senior', 'lead', 'principal']):
            return 7
        elif any(keyword in text_lower for keyword in ['junior', 'entry', 'fresher']):
            return 1
        else:
            return 3
    
    def get_fallback_analysis(self, resume_text):
        """Fallback analysis when models are not available"""
        skills = self.extract_skills_traditional(resume_text)
        experience = self.predict_experience_traditional(resume_text)
        
        return {
            'predicted_category': "Software Engineering",
            'predicted_experience': experience,
            'extracted_skills': skills,
            'match_score': 70.0,
            'analysis_method': 'fallback',
            'model_confidence': 0.5,
            'bert_enhanced': False
        }
    
    def analyze_resume(self, resume_text):
        """Main analyze_resume method for compatibility"""
        return self.analyze_resume_bert(resume_text)
    
    def get_job_recommendations(self, analysis_result, top_k=5):
        """Get job recommendations based on analysis"""
        try:
            category = analysis_result.get('predicted_category', 'Software Engineering')
            skills = analysis_result.get('extracted_skills', [])
            experience = analysis_result.get('predicted_experience', 3)
            
            # Sample job recommendations based on category and skills
            recommendations = []
            
            if 'Data Science' in category or any(skill in ['Python', 'Machine Learning', 'Pandas'] for skill in skills):
                recommendations.extend([
                    {
                        'title': 'Data Scientist',
                        'company': 'Tech Corp',
                        'location': 'Remote',
                        'match_score': 85,
                        'required_skills': ['Python', 'Machine Learning', 'SQL'],
                        'salary_range': '$80k - $120k'
                    },
                    {
                        'title': 'ML Engineer',
                        'company': 'AI Startup',
                        'location': 'San Francisco',
                        'match_score': 82,
                        'required_skills': ['Python', 'TensorFlow', 'AWS'],
                        'salary_range': '$90k - $140k'
                    }
                ])
            
            if 'Frontend' in category or any(skill in ['React', 'JavaScript', 'HTML'] for skill in skills):
                recommendations.extend([
                    {
                        'title': 'Frontend Developer',
                        'company': 'Web Solutions',
                        'location': 'New York',
                        'match_score': 88,
                        'required_skills': ['React', 'JavaScript', 'CSS'],
                        'salary_range': '$70k - $110k'
                    }
                ])
            
            if 'Backend' in category or any(skill in ['Python', 'Java', 'Node.js'] for skill in skills):
                recommendations.extend([
                    {
                        'title': 'Backend Developer',
                        'company': 'Enterprise Solutions',
                        'location': 'Austin',
                        'match_score': 86,
                        'required_skills': ['Python', 'Django', 'PostgreSQL'],
                        'salary_range': '$75k - $115k'
                    }
                ])
            
            # Default recommendations if none match
            if not recommendations:
                recommendations = [
                    {
                        'title': 'Software Engineer',
                        'company': 'Generic Tech',
                        'location': 'Remote',
                        'match_score': 75,
                        'required_skills': ['Programming', 'Problem Solving'],
                        'salary_range': '$65k - $100k'
                    }
                ]
            
            return {
                'recommendations': recommendations[:top_k],
                'total_found': len(recommendations),
                'search_criteria': {
                    'category': category,
                    'skills': skills[:5],  # Top 5 skills
                    'experience_level': experience
                }
            }
            
        except Exception as e:
            print(f"⚠️ Job recommendations error: {e}")
            return {
                'recommendations': [],
                'total_found': 0,
                'error': str(e)
            }
    
    def skill_gap_analysis(self, resume_text, job_description):
        """Analyze skill gaps between resume and job requirements"""
        try:
            print("🔍 Starting BERT-enhanced skill gap analysis...")
            
            # Extract skills from both resume and job description
            resume_skills = self.extract_skills_traditional(resume_text)
            job_skills = self.extract_skills_traditional(job_description)
            
            # Use BERT for semantic similarity if available
            if hasattr(self, 'sentence_model') and self.sentence_model:
                try:
                    # Get embeddings for skills
                    resume_embeddings = self.sentence_model.encode(resume_skills)
                    job_embeddings = self.sentence_model.encode(job_skills)
                    
                    # Calculate semantic similarity
                    from sklearn.metrics.pairwise import cosine_similarity
                    import numpy as np
                    
                    matching_skills = []
                    missing_skills = []
                    
                    for job_skill in job_skills:
                        job_emb = self.sentence_model.encode([job_skill])
                        similarities = cosine_similarity(job_emb, resume_embeddings)
                        max_similarity = np.max(similarities) if len(similarities) > 0 else 0
                        
                        if max_similarity > 0.7:  # High similarity threshold
                            matching_skills.append({
                                'skill': job_skill,
                                'similarity': float(max_similarity),
                                'matched_resume_skill': resume_skills[np.argmax(similarities[0])] if len(similarities[0]) > 0 else job_skill
                            })
                        else:
                            missing_skills.append({
                                'skill': job_skill,
                                'priority': 'high' if max_similarity < 0.3 else 'medium',
                                'similarity': float(max_similarity)
                            })
                    
                    # Calculate overall match score
                    total_job_skills = len(job_skills)
                    matched_count = len(matching_skills)
                    match_percentage = (matched_count / total_job_skills * 100) if total_job_skills > 0 else 0
                    
                    # Generate recommendations
                    recommendations = []
                    for missing in missing_skills:
                        recommendations.append({
                            'type': 'skill_development',
                            'skill': missing['skill'],
                            'priority': missing['priority'],
                            'estimated_learning_time': self._estimate_learning_time(missing['skill']),
                            'resources': self._get_learning_resources(missing['skill'])
                        })
                    
                    return {
                        'analysis_method': 'bert_enhanced',
                        'match_percentage': round(match_percentage, 1),
                        'matching_skills': matching_skills,
                        'missing_skills': missing_skills,
                        'recommendations': recommendations,
                        'total_job_skills': total_job_skills,
                        'total_resume_skills': len(resume_skills),
                        'bert_enhanced': True
                    }
                    
                except Exception as bert_error:
                    print(f"⚠️ BERT analysis failed, using traditional method: {bert_error}")
            
            # Fallback to traditional analysis
            matching_skills = list(set(resume_skills) & set(job_skills))
            missing_skills = list(set(job_skills) - set(resume_skills))
            
            match_percentage = (len(matching_skills) / len(job_skills) * 100) if job_skills else 0
            
            return {
                'analysis_method': 'traditional',
                'match_percentage': round(match_percentage, 1),
                'matching_skills': [{'skill': skill, 'similarity': 1.0} for skill in matching_skills],
                'missing_skills': [{'skill': skill, 'priority': 'medium'} for skill in missing_skills],
                'recommendations': [
                    {
                        'type': 'skill_development',
                        'skill': skill,
                        'priority': 'medium',
                        'estimated_learning_time': self._estimate_learning_time(skill)
                    } for skill in missing_skills[:5]  # Top 5 missing skills
                ],
                'total_job_skills': len(job_skills),
                'total_resume_skills': len(resume_skills),
                'bert_enhanced': False
            }
            
        except Exception as e:
            print(f"❌ Skill gap analysis error: {e}")
            return {
                'analysis_method': 'error',
                'match_percentage': 0,
                'matching_skills': [],
                'missing_skills': [],
                'recommendations': [],
                'error': str(e),
                'bert_enhanced': False
            }
    
    def _estimate_learning_time(self, skill):
        """Estimate learning time for a skill"""
        time_estimates = {
            'python': '2-3 months',
            'javascript': '2-3 months',
            'react': '1-2 months',
            'node.js': '1-2 months',
            'sql': '1 month',
            'git': '2 weeks',
            'docker': '3-4 weeks',
            'kubernetes': '2-3 months',
            'aws': '2-4 months',
            'machine learning': '3-6 months',
            'tensorflow': '2-3 months',
            'pytorch': '2-3 months'
        }
        return time_estimates.get(skill.lower(), '1-2 months')
    
    def _get_learning_resources(self, skill):
        """Get learning resources for a skill"""
        resources = {
            'python': [
                {'type': 'course', 'name': 'Python for Everybody (Coursera)', 'url': 'https://coursera.org'},
                {'type': 'practice', 'name': 'LeetCode Python', 'url': 'https://leetcode.com'},
                {'type': 'documentation', 'name': 'Python Official Docs', 'url': 'https://docs.python.org'}
            ],
            'javascript': [
                {'type': 'course', 'name': 'JavaScript MDN Guide', 'url': 'https://developer.mozilla.org'},
                {'type': 'practice', 'name': 'freeCodeCamp', 'url': 'https://freecodecamp.org'},
                {'type': 'course', 'name': 'JavaScript.info', 'url': 'https://javascript.info'}
            ],
            'react': [
                {'type': 'course', 'name': 'React Official Tutorial', 'url': 'https://reactjs.org'},
                {'type': 'practice', 'name': 'React Challenges', 'url': 'https://codepen.io'},
                {'type': 'course', 'name': 'Scrimba React Course', 'url': 'https://scrimba.com'}
            ]
        }
        
        default_resources = [
            {'type': 'search', 'name': f'{skill} tutorials on YouTube', 'url': 'https://youtube.com'},
            {'type': 'practice', 'name': f'{skill} projects on GitHub', 'url': 'https://github.com'},
            {'type': 'course', 'name': f'{skill} courses on Udemy', 'url': 'https://udemy.com'}
        ]
        
        return resources.get(skill.lower(), default_resources)

# Initialize the BERT-enhanced service
def get_bert_ml_service():
    """Get the BERT-enhanced ML service instance"""
    if not hasattr(get_bert_ml_service, 'instance'):
        get_bert_ml_service.instance = BERTEnhancedMLService()
    return get_bert_ml_service.instance

# Compatibility function for existing code
def get_ml_service():
    """Get ML service (BERT-enhanced by default)"""
    return get_bert_ml_service()
